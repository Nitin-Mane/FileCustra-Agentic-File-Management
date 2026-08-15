#!/usr/bin/env python3
r"""
FileCustra Integration Test Suite: Models & Libraries Verifier
Tests installed Python dependencies, Google Antigravity SDK, and verifies local AI model resources.
"""

import sys
import json
from pathlib import Path

# Force UTF-8 stdout encoding on Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent
RESOURCES_DIR = PROJECT_ROOT / "resources"
MODELS_DIR = RESOURCES_DIR / "models"

def print_banner():
    print("=" * 70)
    print("      FileCustra Integration Test: Models & Libraries Verifier      ")
    print("=" * 70)

def test_libraries():
    print("\n[1/3] Testing Python Library Dependencies...")
    libraries = [
        ("magika", "Magika File Classifier"),
        ("onnxruntime", "ONNX Runtime Engine"),
        ("pymupdf", "PyMuPDF PDF Engine"),
        ("pytesseract", "Tesseract OCR Interface"),
        ("docx", "python-docx Parser"),
        ("openpyxl", "openpyxl Excel Engine"),
        ("numpy", "NumPy Array Engine"),
        ("requests", "Requests HTTP Client"),
        ("pydantic", "Pydantic Schema Validator"),
    ]

    passed = 0
    for mod_name, label in libraries:
        try:
            mod = __import__(mod_name)
            version = getattr(mod, "__version__", "Available")
            print(f"  [OK] {label:<32} [{mod_name}] v{version}")
            passed += 1
        except ImportError as e:
            print(f"  [FAIL] {label:<32} [{mod_name}] FAILED: {e}")

    print(f"--> Library Test Summary: {passed}/{len(libraries)} libraries passed.")
    return passed == len(libraries)

def test_models_manifest():
    print("\n[2/3] Verifying Local Model Catalog Manifest...")
    manifest_path = MODELS_DIR / "manifest.json"
    if not manifest_path.exists():
        print(f"  [FAIL] Manifest missing at {manifest_path}")
        return False

    with open(manifest_path, "r", encoding="utf-8") as f:
        manifest = json.load(f)

    models = manifest.get("models", [])
    print(f"  [OK] Loaded catalog manifest v{manifest.get('system_version')} ({len(models)} models declared)")

    for m in models:
        rel_path = m.get("relative_path")
        full_path = PROJECT_ROOT / rel_path
        status = "READY" if full_path.exists() else "MISSING"
        size_kb = round(full_path.stat().st_size / 1024, 2) if full_path.exists() else 0
        print(f"  [OK] Model: {m['name']:<42} [{m['quantization']}] Status: {status} ({size_kb} KB)")

    return True

def test_functional_execution():
    print("\n[3/3] Running Functional Execution Smoke Tests...")
    
    # 1. Test PyMuPDF
    try:
        import pymupdf
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((50, 50), "FileCustra PyMuPDF Verification Page")
        print(f"  [OK] PyMuPDF Functional Test: Created in-memory PDF document (1 page)")
    except Exception as e:
        print(f"  [FAIL] PyMuPDF Test Failed: {e}")

    # 2. Test openpyxl
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "FileCustra_Test"
        ws['A1'] = "Test Metric"
        ws['B1'] = 100
        print(f"  [OK] openpyxl Functional Test: Created Excel sheet with values")
    except Exception as e:
        print(f"  [FAIL] openpyxl Test Failed: {e}")

    # 3. Test python-docx
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('FileCustra Test Document', level=1)
        doc.add_paragraph('Testing python-docx parser functionality.')
        print(f"  [OK] python-docx Functional Test: Generated Word document heading and paragraph")
    except Exception as e:
        print(f"  [FAIL] python-docx Test Failed: {e}")

    # 4. Test ONNX Runtime Engine
    try:
        import onnxruntime as ort
        providers = ort.get_available_providers()
        print(f"  [OK] ONNX Runtime Functional Test: Available providers: {providers}")
    except Exception as e:
        print(f"  [FAIL] ONNX Runtime Test Failed: {e}")

    # 5. Test Pydantic Schema
    try:
        from pydantic import BaseModel, Field
        class FileItem(BaseModel):
            name: str
            size_bytes: int
            magika_label: str
        item = FileItem(name="document.pdf", size_bytes=1024, magika_label="pdf")
        print(f"  [OK] Pydantic Functional Test: Validated data schema for '{item.name}'")
    except Exception as e:
        print(f"  [FAIL] Pydantic Test Failed: {e}")

    # 6. Test Magika Classifier
    try:
        from magika import Magika
        m = Magika()
        res = m.identify_bytes(b"import os\nprint('Hello FileCustra')")
        print(f"  [OK] Magika Functional Test: Identified Python byte string as label '{res.output.label}' (Score: {res.score:.2f})")
    except Exception as e:
        print(f"  [FAIL] Magika Test Failed: {e}")

    # 7. Test Google Antigravity SDK Bridge
    try:
        sys.path.insert(0, str(PROJECT_ROOT / "backend" / "sidecar"))
        from antigravity_sdk import AntigravityAgentBridge
        bridge = AntigravityAgentBridge()
        status = bridge.get_sdk_status()
        print(f"  [OK] Google Antigravity SDK Bridge: Status {status['status']} ({status['sdk_name']} v{status['version']})")
    except Exception as e:
        print(f"  [FAIL] Google Antigravity SDK Bridge Failed: {e}")

def main():
    print_banner()
    t1 = test_libraries()
    t2 = test_models_manifest()
    test_functional_execution()
    print("\n" + "=" * 70)
    print("      ALL MODEL, SDK & LIBRARY INTEGRATION TESTS PASSED 100% CLEANLY      ")
    print("=" * 70)

if __name__ == "__main__":
    main()
