#!/usr/bin/env python3
"""
Test suite for FileCustra Parser Workers.
"""

import json
import tempfile
import unittest
from pathlib import Path

from parser_workers import (
    ParserWorkerPool,
    ParseResult,
    PDFParser,
    OfficeParser,
    TextParser,
    ImageParser,
    AudioParser,
    VideoParser,
    create_parser_pool,
)


class TestParseResult(unittest.TestCase):
    """Test ParseResult dataclass."""
    
    def test_to_dict(self):
        result = ParseResult(
            file_path="/test/file.pdf",
            file_name="file.pdf",
            parser_type="pdf",
            success=True,
            content_text="Test content",
            metadata={"pages": 5},
            word_count=100,
        )
        
        d = result.to_dict()
        
        self.assertEqual(d["file_path"], "/test/file.pdf")
        self.assertEqual(d["file_name"], "file.pdf")
        self.assertEqual(d["parser_type"], "pdf")
        self.assertTrue(d["success"])
        self.assertEqual(d["content_text"], "Test content")
        self.assertEqual(d["metadata"]["pages"], 5)
        self.assertEqual(d["word_count"], 100)


class TestTextParser(unittest.TestCase):
    """Test text parser."""
    
    def setUp(self):
        self.parser = TextParser()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_parser_type(self):
        self.assertEqual(self.parser.parser_type, "text")
    
    def test_supported_extensions(self):
        self.assertIn("txt", self.parser.supported_extensions)
        self.assertIn("py", self.parser.supported_extensions)
        self.assertIn("json", self.parser.supported_extensions)
        self.assertIn("md", self.parser.supported_extensions)
    
    def test_can_parse(self):
        self.assertTrue(self.parser.can_parse("test.txt"))
        self.assertTrue(self.parser.can_parse("script.py"))
        self.assertFalse(self.parser.can_parse("image.png"))
    
    def test_parse_text_file(self):
        txt_path = Path(self.temp_dir) / "test.txt"
        txt_path.write_text("Hello world\nThis is a test file\nWith multiple lines")
        
        result = self.parser.parse(str(txt_path))
        
        self.assertTrue(result.success)
        self.assertEqual(result.parser_type, "text")
        self.assertIn("Hello world", result.content_text)
        self.assertEqual(result.metadata["line_count"], 3)
    
    def test_parse_nonexistent_file(self):
        result = self.parser.parse("/nonexistent/file.txt")
        self.assertFalse(result.success)
        self.assertIn("not found", result.error_message.lower())


class TestPDFParser(unittest.TestCase):
    """Test PDF parser."""
    
    def setUp(self):
        self.parser = PDFParser()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_parser_type(self):
        self.assertEqual(self.parser.parser_type, "pdf")
    
    def test_can_parse(self):
        self.assertTrue(self.parser.can_parse("document.pdf"))
        self.assertFalse(self.parser.can_parse("document.txt"))
    
    def test_parse_nonexistent_file(self):
        result = self.parser.parse("/nonexistent/file.pdf")
        self.assertFalse(result.success)


class TestOfficeParser(unittest.TestCase):
    """Test Office parser."""
    
    def setUp(self):
        self.parser = OfficeParser()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_parser_type(self):
        self.assertEqual(self.parser.parser_type, "office")
    
    def test_can_parse(self):
        self.assertTrue(self.parser.can_parse("document.docx"))
        self.assertTrue(self.parser.can_parse("spreadsheet.xlsx"))
        self.assertTrue(self.parser.can_parse("presentation.pptx"))
        self.assertFalse(self.parser.can_parse("document.pdf"))
    
    def test_parse_docx(self):
        try:
            from docx import Document
            
            doc_path = Path(self.temp_dir) / "test.docx"
            doc = Document()
            doc.add_paragraph("Hello world")
            doc.add_paragraph("This is a test document")
            doc.save(str(doc_path))
            
            result = self.parser.parse(str(doc_path))
            
            self.assertTrue(result.success)
            self.assertIn("Hello world", result.content_text)
            self.assertIn("word_count", result.to_dict())
        except ImportError:
            self.skipTest("python-docx not installed")
    
    def test_parse_xlsx(self):
        try:
            from openpyxl import Workbook
            
            xlsx_path = Path(self.temp_dir) / "test.xlsx"
            wb = Workbook()
            ws = wb.active
            ws["A1"] = "Name"
            ws["B1"] = "Value"
            ws["A2"] = "Test"
            ws["B2"] = 123
            wb.save(str(xlsx_path))
            
            result = self.parser.parse(str(xlsx_path))
            
            self.assertTrue(result.success)
            self.assertIn("Name", result.content_text)
            self.assertIn("sheets", result.metadata)
        except ImportError:
            self.skipTest("openpyxl not installed")


class TestImageParser(unittest.TestCase):
    """Test image parser."""
    
    def setUp(self):
        self.parser = ImageParser()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_parser_type(self):
        self.assertEqual(self.parser.parser_type, "image")
    
    def test_can_parse(self):
        self.assertTrue(self.parser.can_parse("image.png"))
        self.assertTrue(self.parser.can_parse("photo.jpg"))
        self.assertFalse(self.parser.can_parse("document.txt"))
    
    def test_parse_image(self):
        try:
            from PIL import Image
            
            img_path = Path(self.temp_dir) / "test.png"
            img = Image.new("RGB", (100, 100), color="red")
            img.save(str(img_path))
            
            result = self.parser.parse(str(img_path))
            
            self.assertTrue(result.success)
            self.assertEqual(result.metadata["width"], 100)
            self.assertEqual(result.metadata["height"], 100)
        except ImportError:
            self.skipTest("Pillow not installed")


class TestParserWorkerPool(unittest.TestCase):
    """Test parser worker pool."""
    
    def setUp(self):
        self.pool = create_parser_pool()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_create_pool(self):
        pool = create_parser_pool()
        self.assertIsInstance(pool, ParserWorkerPool)
    
    def test_get_parser_types(self):
        types = self.pool.get_parser_types()
        self.assertIn("text", types)
        self.assertIn("pdf", types)
        self.assertIn("office", types)
        self.assertIn("image", types)
    
    def test_get_supported_extensions(self):
        extensions = self.pool.get_supported_extensions()
        self.assertIn("txt", extensions)
        self.assertIn("pdf", extensions)
        self.assertIn("docx", extensions)
        self.assertIn("png", extensions)
    
    def test_parse_text_file(self):
        txt_path = Path(self.temp_dir) / "test.txt"
        txt_path.write_text("Hello world")
        
        result = self.pool.parse_file(str(txt_path))
        
        self.assertTrue(result.success)
        self.assertEqual(result.parser_type, "text")
    
    def test_parse_nonexistent_file(self):
        result = self.pool.parse_file("/nonexistent/file.txt")
        self.assertFalse(result.success)
    
    def test_parse_batch(self):
        files = []
        for name in ["test1.txt", "test2.txt"]:
            path = Path(self.temp_dir) / name
            path.write_text(f"Content of {name}")
            files.append(str(path))
        
        results = self.pool.parse_batch(files)
        
        self.assertEqual(len(results), 2)
        self.assertTrue(all(r.success for r in results))


class TestParserIntegration(unittest.TestCase):
    """Integration tests for parsers."""
    
    def setUp(self):
        self.pool = create_parser_pool()
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_json_parsing(self):
        json_path = Path(self.temp_dir) / "data.json"
        json_path.write_text(json.dumps({"key": "value", "number": 42}))
        
        result = self.pool.parse_file(str(json_path))
        
        self.assertTrue(result.success)
        self.assertIn("key", result.content_text)
    
    def test_markdown_parsing(self):
        md_path = Path(self.temp_dir) / "readme.md"
        md_path.write_text("# Title\n\nThis is a **bold** word.")
        
        result = self.pool.parse_file(str(md_path))
        
        self.assertTrue(result.success)
        self.assertIn("Title", result.content_text)
    
    def test_python_parsing(self):
        py_path = Path(self.temp_dir) / "script.py"
        py_path.write_text("#!/usr/bin/env python3\nimport os\nprint('hello')")
        
        result = self.pool.parse_file(str(py_path))
        
        self.assertTrue(result.success)
        self.assertIn("import os", result.content_text)


if __name__ == "__main__":
    unittest.main()
